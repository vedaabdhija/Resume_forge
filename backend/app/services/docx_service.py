import os
from docx import Document
from typing import Dict, Any
from app.core.logging import logger

class DocxService:
    @classmethod
    def tailor_docx_in_place(
        cls,
        original_docx_path: str,
        output_docx_path: str,
        original_json: Dict[str, Any],
        tailored_json: Dict[str, Any]
    ) -> bool:
        """
        Loads the original DOCX resume and updates paragraph text in-place.
        Preserves fonts, sizing, colors, and margins by modifying existing runs.
        """
        if not os.path.exists(original_docx_path):
            logger.error(f"Original DOCX file not found: {original_docx_path}")
            return False
            
        try:
            doc = Document(original_docx_path)
            
            # Map original items to tailored items for direct replacement
            replacements = cls._build_replacements_map(original_json, tailored_json)
            
            # 1. Search and replace in paragraphs
            for p in doc.paragraphs:
                cls._replace_in_paragraph(p, replacements)
                
            # 2. Search and replace inside tables (common in resume templates)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            cls._replace_in_paragraph(p, replacements)
                            
            doc.save(output_docx_path)
            logger.info(f"DOCX tailored in-place successfully: {output_docx_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to tailor DOCX in-place: {str(e)}")
            return False

    @staticmethod
    def _build_replacements_map(original: Dict[str, Any], tailored: Dict[str, Any]) -> Dict[str, str]:
        """Maps original bullet points and lists to tailored outputs."""
        map_dict = {}
        
        # Map experience highlights
        orig_exp = original.get("experience", [])
        tail_exp = tailored.get("experience", [])
        
        for o_job in orig_exp:
            o_role = o_job.get("role", "")
            o_bullets = o_job.get("highlights", [])
            
            # Find matching tailored job by role/company
            t_job = next((j for j in tail_exp if j.get("role") == o_role), None)
            if t_job:
                t_bullets = t_job.get("highlights", [])
                for idx, o_bullet in enumerate(o_bullets):
                    if idx < len(t_bullets):
                        map_dict[o_bullet] = t_bullets[idx]
                        
        # Map project highlights
        orig_proj = original.get("projects", [])
        tail_proj = tailored.get("projects", [])
        for o_p in orig_proj:
            o_title = o_p.get("title", "")
            o_bullets = o_p.get("highlights", [])
            
            t_p = next((p for p in tail_proj if p.get("title") == o_title), None)
            if t_p:
                t_bullets = t_p.get("highlights", [])
                for idx, o_bullet in enumerate(o_bullets):
                    if idx < len(t_bullets):
                        map_dict[o_bullet] = t_bullets[idx]
                        
        # Map skills list string (if exists as a comma list)
        o_skills = original.get("skills", [])
        t_skills = tailored.get("skills", [])
        if o_skills and t_skills:
            map_dict[", ".join(o_skills)] = ", ".join(t_skills)
            
        return map_dict

    @staticmethod
    def _replace_in_paragraph(p, replacements: Dict[str, str]):
        """Replaces text in a paragraph runs without breaking style attributes."""
        if not p.runs:
            return
            
        # Get full paragraph text to check for matches
        full_text = "".join(run.text for run in p.runs)
        
        for old_txt, new_txt in replacements.items():
            # If the paragraph matches or contains the old bullet point
            if old_txt in full_text or full_text in old_txt:
                # Write the new text to the first run
                p.runs[0].text = new_txt
                # Clear all other runs to prevent duplicate text
                for run in p.runs[1:]:
                    run.text = ""
                break
